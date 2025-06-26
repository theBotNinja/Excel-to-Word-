import pandas as pd
from docx import Document
import sys
from docx.shared import Inches
import os
from time import sleep
from tkinter import filedialog

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def SelectCol(df):
    col_list =  df.columns.tolist()
    col_selected_index = 0
    while col_selected_index < len(col_list):
        clear()
        print(f"{bcolors.BOLD} {bcolors.OKBLUE} {bcolors.UNDERLINE}Select a column first {bcolors.ENDC}\n")
        print(f"{bcolors.BOLD} Showing - {bcolors.ENDC} {bcolors.HEADER} {col_list[col_selected_index]}{bcolors.ENDC}\n")
        print(df[col_list[col_selected_index]],end="\n\n")
        ans = input(f"{bcolors.BOLD}{bcolors.OKGREEN}Press y to select the column {bcolors.ENDC} or {bcolors.BOLD}{bcolors.OKBLUE} Enter to show next column (b for back) : {bcolors.ENDC}")
        if ans == "y":
            return col_list[col_selected_index]
        if ans == "b":
            col_selected_index-=1
            continue
        col_selected_index+=1
    return None

def SelectRow(df,col):
    clear()
    print(f"{bcolors.BOLD} {bcolors.OKBLUE} {bcolors.UNDERLINE}Select a row index to start from (default 0){bcolors.ENDC}\n")
    print(df[col])
    while 1:
        print()
        ans = int(input(f"{bcolors.BOLD}{bcolors.OKBLUE}Enter the index of row ({bcolors.FAIL}press -1 to exit{bcolors.OKBLUE}) : {bcolors.ENDC}"))
        if ans in  range(len(df[col])):
            return ans
        if ans  == -1:
            return 0
    return 0

def create_questionnaire(questions, options, heading="Survey Form"):
    doc = Document()
    # Add a title heading
    doc.add_heading(heading, level=1)
    # Add questions and options
    for i, question in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {question}")
        char = 65
        for option in options:
            p = doc.add_paragraph(f"{chr(char)}. {option}")
            p.paragraph_format.left_indent = Inches(0.5)
            char+=1

    # Save document
    doc.save("./word_files/"+heading+".docx")
    print(f"{bcolors.OKGREEN}Document saved as {heading}.docx 👍{bcolors.ENDC}")

def clear():
    os.system('cls' if os.name == 'nt' else 'clear')

def SelectOptions():
    options_list = {
        "3-Point Scale":["Agree","Neutral","Disagree"],
        "5-Point Scale":["Strongly Agree","Agree","Neutral","Disagree","Strongly Disagree"],
        "Frequency Scale":["Always","Often","Sometimes","Rarely","Never"],
        "Importance Scale":["Very Important","Important","Neutral","Unimportant","Very Unimportant"],
        "Satisfaction Scale":["Very Satisfied","Satisfied","Neutral","Dissatisfied","Very Dissatisfied"]
    }
    lis = list(enumerate(options_list.keys()))
    while 1:
        clear()
        print(f"{bcolors.BOLD} {bcolors.OKBLUE} {bcolors.UNDERLINE}Select option type {bcolors.ENDC}\n")
        for item in lis:
            print(f"{bcolors.OKBLUE}{item[0]}{bcolors.ENDC+bcolors.BOLD} -> {bcolors.ENDC}{item[1]}")
        print()
        ans = int(input(f"{bcolors.BOLD+bcolors.OKBLUE}Enter the index : {bcolors.ENDC}"))
        if ans in range(len(options_list.keys())):
            clear()
            print(f"{bcolors.BOLD} {bcolors.OKBLUE} {bcolors.UNDERLINE}Selected option will look like this : {bcolors.ENDC}\n")
            for option in options_list[lis[ans][1]]:
                print(f" ◉ {option}")
            print()
            if "y"==input(f"{bcolors.BOLD+bcolors.OKGREEN} Press y for final select {bcolors.ENDC} and {bcolors.BOLD+bcolors.OKBLUE} Enter to restart {bcolors.ENDC} : "):
                return options_list[lis[ans][1]]

print(f"""{bcolors.HEADER+bcolors.BOLD} Excel -> Word {bcolors.ENDC} for forms
This script converts excels to word that can be then used for creating quick forms.
Start by giving a excel file to the script ...
""")

# ask for excel file 
file = None
if len(sys.argv) > 1:
    file = sys.argv[1]
else:
    file = filedialog.askopenfilename()
if file=="":
    print(f"{bcolors.BOLD+bcolors.FAIL} -X-  No file path was provided  -X- {bcolors.ENDC}")
    exit()
try:
    os.mkdir("word_files")
except:
    print(f"{bcolors.WARNING} word_files alraedy exits {bcolors.BOLD}(might cause over writes){bcolors.ENDC}\n")
    sleep(2)
# reading excel file
print(f"{bcolors.BOLD+bcolors.OKBLUE}Loading file ... {bcolors.ENDC}")
ef = pd.ExcelFile(file)
sheets = ef.sheet_names
ef.close()
col_selected = 0
row_start = 0
df = pd.read_excel("renu.xlsx",sheet_name=sheets[0]) #first sheet for preview 
OneByOne = input(f"{bcolors.BOLD+bcolors.OKGREEN}file load successfull! {bcolors.ENDC} \n\n{bcolors.OKBLUE}enter to continue {bcolors.WARNING}or{bcolors.OKCYAN} press y for One-by-One mode{bcolors.ENDC} : ")
if OneByOne=="":
    # Selection of row and column
    col_selected = SelectCol(df)
    if not col_selected :
        print(f"{bcolors.BOLD+bcolors.FAIL} -X-  No Column Selected  -X- {bcolors.ENDC}")
        exit()
    row_start = SelectRow(df,col_selected)
    option_list = SelectOptions()
    clear()
    print(f"{bcolors.OKBLUE} Selected column is - {bcolors.ENDC}{col_selected}")
    print(f"{bcolors.OKBLUE} Selected row index is - {bcolors.ENDC}{row_start}\n")
    print(f"{bcolors.WARNING} starting .... {bcolors.ENDC}")

    ## Selection done

    for i in sheets:
        df = pd.read_excel("renu.xlsx",sheet_name=i)
        Question_list = df[col_selected][row_start:].dropna().tolist()
        create_questionnaire(Question_list,option_list,i)
elif OneByOne == "about":
    print(f"made by - {bcolors.BOLD+bcolors.HEADER}@theBotNinja{bcolors.ENDC}")
else:
    option_list = SelectOptions()
    for i in sheets:
        col_selected = SelectCol(df)
        if not col_selected :
            print(f"{bcolors.BOLD+bcolors.FAIL} -X-  No Column Selected  -X- {bcolors.ENDC}")
            exit()
        row_start = SelectRow(df,col_selected)
        df = pd.read_excel("renu.xlsx",sheet_name=i)
        clear()
        print(f"{bcolors.OKBLUE} Selected column is - {bcolors.ENDC}{col_selected}")
        print(f"{bcolors.OKBLUE} Selected row index is - {bcolors.ENDC}{row_start}\n")
        print(f"{bcolors.WARNING} starting .... {bcolors.ENDC}")
        Question_list = df[col_selected][row_start:].dropna().tolist()
        create_questionnaire(Question_list,option_list,i)